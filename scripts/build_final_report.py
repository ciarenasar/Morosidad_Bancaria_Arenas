"""Construye el informe académico final del proyecto de morosidad bancaria."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "informe_final_morosidad_bancaria.docx"
FIGURES = ROOT / "reports" / "figures"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(82, 92, 104)
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)
GREEN = RGBColor(36, 94, 64)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError("Los anchos de tabla deben sumar 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float | None = None, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "es-CL")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, *, size, color, before, after, line, bold=False, alignment=None) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if alignment is not None:
        pf.alignment = alignment


def add_custom_numbering(doc: Document, *, abstract_id: int, num_id: int, bullet: bool) -> None:
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)


def add_numbering_instance(doc: Document, *, abstract_id: int, num_id: int) -> None:
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    if abstract_id == 901 and num_id != 901:
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        num.append(lvl_override)
    numbering.append(num)


def add_list_item(doc: Document, text: str, *, numbered=False, restart=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    if numbered:
        if restart or not hasattr(doc, "_report_num_id"):
            next_id = getattr(doc, "_report_num_id", 901) + 1
            setattr(doc, "_report_num_id", next_id)
            add_numbering_instance(doc, abstract_id=901, num_id=next_id)
        list_num_id = getattr(doc, "_report_num_id")
    else:
        list_num_id = 900
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(list_num_id))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    set_run_font(p.add_run(text), size=11)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=11, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=11)
    else:
        set_run_font(p.add_run(text), size=11)


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Table Citation")
    set_run_font(p.add_run(f"Fuente: {text}"), size=8.5, color=GRAY, italic=True)


def add_callout(doc: Document, title: str, text: str, *, tone="blue") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    fill = PALE_BLUE if tone == "blue" else ("FFF4CE" if tone == "gold" else "FDECEC")
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    color = DARK_BLUE if tone == "blue" else (GOLD if tone == "gold" else RED)
    set_run_font(p.add_run(title), size=11, color=color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run_font(p2.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], source: str) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=9, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(value)), size=8.8)
    set_table_geometry(table, widths)
    add_source(doc, source)


def add_figure(doc: Document, filename: str, caption: str, source: str, *, width=6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    cap.paragraph_format.keep_with_next = True
    set_run_font(cap.add_run(caption), size=9.5, bold=True, color=NAVY)
    add_source(doc, source)


def add_heading(doc: Document, text: str, level=1, *, page_break=False) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    configure_style(styles["Normal"], size=11, color=RGBColor(0, 0, 0), before=0, after=6, line=1.10)
    configure_style(styles["Heading 1"], size=16, color=BLUE, before=16, after=8, line=1.0, bold=True)
    configure_style(styles["Heading 2"], size=13, color=BLUE, before=12, after=6, line=1.0, bold=True)
    configure_style(styles["Heading 3"], size=12, color=DARK_BLUE, before=8, after=4, line=1.0, bold=True)
    styles.add_style("Table Citation", 1)
    configure_style(styles["Table Citation"], size=8.5, color=GRAY, before=4, after=4, line=1.0)
    styles.add_style("Cover Kicker", 1)
    configure_style(styles["Cover Kicker"], size=11, color=GOLD, before=0, after=18, line=1.0, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_custom_numbering(doc, abstract_id=900, num_id=900, bullet=True)
    add_custom_numbering(doc, abstract_id=901, num_id=901, bullet=False)
    add_numbering_instance(doc, abstract_id=900, num_id=900)
    add_numbering_instance(doc, abstract_id=901, num_id=901)
    doc.settings.odd_and_even_pages_header_footer = False

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("Proyecto del curso de Finanzas  |  Página "), size=8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)

    props = doc.core_properties
    props.title = "Pronóstico de morosidad bancaria en Chile"
    props.subject = "Informe final del proyecto de machine learning financiero"
    props.author = "Proyecto del curso de Finanzas"
    props.keywords = "morosidad, Chile, CMF, Banco Central, machine learning, point-in-time"


def build_report() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Portada editorial.
    doc.add_paragraph().paragraph_format.space_after = Pt(86)
    doc.add_paragraph("PROYECTO FINAL · FINANZAS", style="Cover Kicker")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run("Pronóstico de morosidad\nbancaria en Chile"), size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_run_font(
        p.add_run("Evaluación point-in-time de modelos de regresión y alerta de estrés"),
        size=15,
        color=DARK_BLUE,
    )
    add_callout(
        doc,
        "Conclusión principal",
        "En el horizonte de seis meses, ningún modelo aprendido mejora de forma robusta al benchmark de cambio cero. La señal de estrés aporta ranking, pero no probabilidades calibradas ni una alerta lista para producción.",
        tone="blue",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Muestra mensual: marzo de 2014 – junio de 2026"), size=11, color=GRAY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Cierre metodológico y técnico · 17 de agosto de 2026"), size=10, color=GRAY)

    # Resumen ejecutivo.
    add_heading(doc, "Resumen ejecutivo", page_break=True)
    add_body(
        doc,
        "Este proyecto investiga si la información macrofinanciera disponible en tiempo real permite anticipar el cambio a seis meses de la morosidad de 90 días o más de la cartera de consumo del sistema bancario chileno. La unidad es el sistema agregado y la fecha de emisión coincide con la publicación efectiva de la información CMF de cada mes."
    )
    add_body(
        doc,
        "El diseño reproduce la información que habría tenido un analista en cada fecha: cada predictor debe cumplir available_date ≤ fecha de emisión, las etiquetas todavía no publicadas se purgan y el holdout de enero de 2024 a diciembre de 2025 permanece completamente cerrado. La base contiene 148 meses continuos, ocho series macroeconómicas y 23 variables modelables."
    )
    add_callout(
        doc,
        "Resultado de decisión",
        "Cambio cero se mantiene como campeón h=6 (MAE 0,434 pp). ElasticNet es el mejor challenger aprendido (MAE 0,507 pp), pero empeora el MAE en 16,8%. La regresión logística lidera el ranking de estrés (AP 0,771), aunque su Brier 0,277 es peor que el 0,248 de la prevalencia histórica.",
        tone="gold",
    )
    add_heading(doc, "Síntesis de la evidencia", level=2)
    for item in [
        "Los 46 orígenes externos abarcan marzo de 2020 a diciembre de 2023; ningún registro del holdout participa en ajuste, selección o evaluación.",
        "El bootstrap por bloques no prueba superioridad de ElasticNet en h=6 y los drivers rotan considerablemente entre bloques.",
        "A doce meses aparece una mejora puntual de 2,5% con ventana expansiva, pero el intervalo de incertidumbre incluye cero y el resultado es inestable.",
        "La alerta logística separa razonablemente los meses de estrés, pero reacciona con seis meses de retraso al principal episodio y omite el evento aislado de agosto de 2023.",
        "El proyecto es reproducible desde archivos raw locales: 11 de 11 controles técnicos pasan y no existen violaciones de disponibilidad ni aperturas del holdout.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "Contenido", level=2)
    add_table(
        doc,
        ["Sección", "Pregunta que responde"],
        [
            ["1–3", "Qué se pronostica, con qué datos y bajo qué protocolo temporal"],
            ["4–7", "Qué tan bien funcionan benchmarks, modelos ML y sus explicaciones"],
            ["8–10", "Qué cambia por horizonte, estrés y régimen macrofinanciero"],
            ["11–13", "Qué se concluye, cuáles son los límites y cómo se reproduce"],
        ],
        [1800, 7560],
        "Elaboración propia.",
    )

    # 1.
    add_heading(doc, "1. Problema, alcance y pregunta de investigación", page_break=True)
    add_body(
        doc,
        "La morosidad bancaria responde con rezagos y posibles no linealidades a la actividad, las tasas, la inflación, el empleo y la dinámica del crédito. El desafío no es solo predictivo: también es temporal. Una evaluación válida debe excluir cualquier dato que todavía no estuviera publicado cuando se habría emitido el pronóstico."
    )
    add_heading(doc, "Pregunta", level=2)
    add_callout(
        doc,
        "Pregunta central",
        "¿Puede un conjunto pequeño y auditable de señales macrofinancieras mejorar el pronóstico a seis meses del cambio en la morosidad de consumo, frente a reglas simples que no requieren estimación?",
    )
    add_heading(doc, "Definición del MVP", level=2)
    add_table(
        doc,
        ["Elemento", "Definición"],
        [
            ["Unidad", "Sistema bancario chileno agregado"],
            ["Cartera", "Consumo"],
            ["Frecuencia", "Mensual"],
            ["Variable base", "Razón de mora de 90 días o más (%)"],
            ["Target", "mora(t+6) − mora(t), en puntos porcentuales"],
            ["Fecha de emisión", "Publicación CMF efectiva del mes t"],
            ["Evaluación", "Walk-forward temporal, sin barajar observaciones"],
        ],
        [2000, 7360],
        "docs/project_charter.md y ADR-001 a ADR-004.",
    )
    add_heading(doc, "Hipótesis de trabajo", level=2)
    for index, item in enumerate([
        "La historia reciente de mora y las condiciones macroeconómicas contienen señal incremental sobre el cambio futuro.",
        "La ventaja, si existe, debe persistir frente a un benchmark de cambio cero y bajo remuestreo por bloques que reconozca la dependencia de errores solapados.",
        "La utilidad potencial de una alerta de estrés depende tanto del ranking como de la calibración y del costo de las falsas alarmas.",
    ]):
        add_list_item(doc, item, numbered=True, restart=index == 0)

    # 2.
    add_heading(doc, "2. Datos e información point-in-time", page_break=True)
    add_body(
        doc,
        "La variable objetivo proviene de APIBEST de la Comisión para el Mercado Financiero. Las variables macroeconómicas provienen de la Base de Datos Estadísticos del Banco Central de Chile. La integración mensual conserva observation_date y available_date para cada serie."
    )
    add_table(
        doc,
        ["Control", "Resultado"],
        [
            ["Cobertura CMF", "148/148 meses; marzo 2014–junio 2026"],
            ["Calendario efectivo CMF", "148/148 meses; rezago 25–38 días"],
            ["Series BCCh", "8 series; 148/148 meses cada una"],
            ["Predictores faltantes tras corte as-of", "0"],
            ["Violaciones de available_date", "0"],
            ["Variables finales", "23"],
            ["Filas completas de desarrollo", "106"],
            ["Holdout reservado", "24 meses; enero 2024–diciembre 2025"],
        ],
        [3300, 6060],
        "docs/data_coverage.md, docs/joint_coverage.md y technical_acceptance_v001.json.",
    )
    add_heading(doc, "Predictores macroeconómicos", level=2)
    for item in [
        "IMACEC desestacionalizado e inflación anual IPC.",
        "Desempleo nacional, tasa de política monetaria y tipo de cambio observado.",
        "Tasa de créditos de consumo, colocaciones reales de consumo y M1 real.",
        "Historia de mora mediante nivel, rezagos, cambios, medias móviles y estacionalidad mensual.",
    ]:
        add_list_item(doc, item)
    add_callout(
        doc,
        "Limitación de vintage",
        "El corte as-of respeta las fechas de publicación, pero las series BCCh son el vintage vigente descargado. Las revisiones históricas pueden introducir sesgo de vintage aunque no exista look-ahead por publicación.",
        tone="gold",
    )

    # 3.
    add_heading(doc, "3. Diseño experimental y gobierno del holdout", page_break=True)
    add_body(
        doc,
        "El backtest externo utiliza 46 orígenes mensuales entre marzo de 2020 y diciembre de 2023. En cada origen se entrena solo con etiquetas cuya publicación ya ocurrió. La muestra efectiva aumenta de 55 a 100 observaciones en h=6 y se purgan como máximo cinco etiquetas."
    )
    add_heading(doc, "Secuencia de evaluación", level=2)
    for index, item in enumerate([
        "Reconstruir la fecha de emisión desde la publicación CMF efectiva.",
        "Aplicar el corte de disponibilidad a cada predictor y construir las 23 variables sin usar información futura.",
        "Purgar targets t+6 todavía no publicados en la fecha de entrenamiento.",
        "Seleccionar hiperparámetros dentro de cada bloque externo mediante validación temporal anidada y purgada.",
        "Registrar la predicción externa, las métricas y la trazabilidad del origen.",
        "Comparar diferencias de pérdida con bootstrap móvil circular de seis meses.",
    ]):
        add_list_item(doc, item, numbered=True, restart=index == 0)
    add_callout(
        doc,
        "Regla irreversible de evaluación",
        "El holdout no se usa para elegir variables, modelos, hiperparámetros, umbrales, ventanas ni conclusiones del informe. Su apertura requiere un protocolo posterior y una decisión explícita.",
        tone="red",
    )

    # 4.
    add_heading(doc, "4. Benchmarks: la regla simple fija una vara exigente", page_break=True)
    add_body(
        doc,
        "Cinco referencias simples cuantifican cuánto valor agrega realmente un modelo entrenado. El benchmark de cambio cero pronostica que la razón de mora será igual a la observada en t cuando se evalúe t+6."
    )
    add_table(
        doc,
        ["Benchmark", "MAE", "RMSE", "Dirección", "Mejora vs. cero"],
        [
            ["Cambio cero", "0,434", "0,566", "0,0%", "0,0%"],
            ["Último cambio conocido", "0,554", "0,712", "67,4%", "−27,7%"],
            ["Promedio 12 cambios", "0,618", "0,727", "54,3%", "−42,5%"],
            ["Ingenuo estacional", "0,730", "0,822", "43,5%", "−68,3%"],
            ["OLS autorregresiva", "0,628", "0,795", "60,9%", "−44,8%"],
        ],
        [3000, 1300, 1300, 1700, 2060],
        "reports/tables/baseline_metrics_development.csv.",
    )
    add_body(
        doc,
        "Cambio cero domina en MAE, RMSE y error absoluto mediano. El último cambio conocido anticipa mejor el signo, pero su mayor error de magnitud muestra que exactitud direccional y precisión puntual responden a objetivos distintos."
    )

    # 5.
    add_heading(doc, "5. Modelos aprendidos: ElasticNet lidera, pero no vence", page_break=True)
    add_body(
        doc,
        "Se compararon ElasticNet, Random Forest y XGBoost con selección interna temporal. Los tres modelos usan el mismo feature set y los mismos orígenes que los benchmarks. En total se evaluaron 112 combinaciones modelo-bloque y 1.344 predicciones internas."
    )
    add_table(
        doc,
        ["Modelo", "MAE", "RMSE", "Dirección", "Mejora vs. cero"],
        [
            ["Cambio cero", "0,434", "0,566", "0,0%", "0,0%"],
            ["ElasticNet", "0,507", "0,635", "41,3%", "−16,8%"],
            ["XGBoost", "0,537", "0,677", "41,3%", "−23,8%"],
            ["Random Forest", "0,584", "0,734", "52,2%", "−34,6%"],
        ],
        [3100, 1300, 1300, 1650, 2010],
        "reports/tables/model_comparison_development.csv.",
    )
    add_figure(
        doc,
        "oos_predictions_timeline.png",
        "Figura 1. Pronósticos externos y cambio observado a seis meses",
        "reports/figures/oos_predictions_timeline.png; predicciones de desarrollo.",
        width=6.4,
    )
    add_body(
        doc,
        "La mayor dificultad aparece en abril–agosto de 2020: el cambio observado cae entre 0,84 y 1,40 pp, mientras los modelos anticipan aumentos pequeños. En el tramo final, con variaciones reales menores, los árboles también producen errores amplificados."
    )

    # 6.
    add_heading(doc, "6. Parsimonia e incertidumbre: la desventaja no desaparece", page_break=True)
    add_body(
        doc,
        "Reducir ElasticNet a núcleos autorregresivo, macro o mixto no corrige su desempeño. La versión autorregresiva de cinco variables se acerca al modelo completo, pero sigue por debajo del benchmark."
    )
    add_table(
        doc,
        ["Especificación", "Variables", "MAE", "Mejora vs. cero"],
        [
            ["ElasticNet completo", "23", "0,507", "−16,8%"],
            ["Núcleo autorregresivo", "5", "0,519", "−19,6%"],
            ["Núcleo macro", "8", "0,664", "−53,1%"],
            ["Núcleo mixto", "10", "0,682", "−57,2%"],
        ],
        [4200, 1300, 1600, 2260],
        "reports/tables/parsimonious_model_comparison.csv.",
    )
    add_heading(doc, "Bootstrap por bloques", level=2)
    add_body(
        doc,
        "Para ElasticNet completo, la diferencia de MAE frente a cambio cero es +0,073 pp, con IC 95% [−0,019; +0,186]. El intervalo incluye cero: no se declara una diferencia estadísticamente concluyente. Aun así, el estimador puntual es adverso, solo gana 28,3% de los meses y la probabilidad bootstrap de mejora es 6,6%."
    )
    add_callout(
        doc,
        "Lectura correcta",
        "La falta de significancia no demuestra equivalencia. Con una muestra corta y errores solapados, la decisión prudente combina el intervalo con el signo, la magnitud y la estabilidad temporal de la diferencia.",
        tone="gold",
    )

    # 7.
    add_heading(doc, "7. Explicabilidad: drivers reconocibles, ranking inestable", page_break=True)
    add_body(
        doc,
        "Las explicaciones se recalcularon para cada origen externo con su muestra purgada. Se obtuvieron 3.174 atribuciones y la reconstrucción aditiva difiere de la predicción registrada en menos de 2,11×10⁻⁷."
    )
    add_figure(
        doc,
        "global_feature_importance.png",
        "Figura 2. Drivers globales fuera de muestra por familia de modelo",
        "reports/figures/global_feature_importance.png.",
        width=6.4,
    )
    add_body(
        doc,
        "ElasticNet se apoya en IMACEC anual (30,2%), media de mora a 12 meses (20,2%) y colocaciones reales de consumo (11,1%). XGBoost y Random Forest priorizan la historia de morosidad. Las medidas de importancia no son directamente comparables entre familias."
    )
    add_heading(doc, "Estabilidad entre bloques", level=2)
    add_table(
        doc,
        ["Modelo", "Spearman medio", "Mínimo", "Máximo"],
        [
            ["ElasticNet", "0,21", "−0,16", "0,50"],
            ["XGBoost", "0,41", "0,07", "0,68"],
            ["Random Forest", "0,54", "0,14", "0,85"],
        ],
        [3600, 2100, 1830, 1830],
        "reports/tables/fold_rank_correlations.csv.",
    )
    add_body(
        doc,
        "La rotación de ElasticNet y el episodio de junio de 2021 —cuando un fuerte efecto base del IMACEC induce una sobreestimación— aconsejan leer los drivers como explicaciones del algoritmo, no como evidencia causal."
    )

    # 8.
    add_heading(doc, "8. Robustez por horizonte y ventana", page_break=True)
    add_figure(
        doc,
        "horizon_robustness.png",
        "Figura 3. Error y aporte de ElasticNet por horizonte y ventana",
        "reports/figures/horizon_robustness.png.",
        width=6.4,
    )
    add_table(
        doc,
        ["Horizonte", "Ventana", "MAE cero", "MAE EN", "Mejora", "Dirección EN"],
        [
            ["3", "Expansiva", "0,246", "0,287", "−16,7%", "26,1%"],
            ["3", "Móvil", "0,246", "0,296", "−20,5%", "30,4%"],
            ["6", "Expansiva", "0,434", "0,507", "−16,8%", "41,3%"],
            ["6", "Móvil", "0,434", "0,522", "−20,4%", "37,0%"],
            ["12", "Expansiva", "0,660", "0,643", "+2,5%", "82,6%"],
            ["12", "Móvil", "0,660", "0,685", "−3,9%", "78,3%"],
        ],
        [1200, 1800, 1500, 1500, 1500, 1860],
        "reports/tables/horizon_robustness_metrics.csv.",
    )
    add_body(
        doc,
        "El caso h=12 expansivo es la única mejora puntual (+2,5%), con dirección correcta en 82,6% de los orígenes. Sin embargo, su diferencia de MAE tiene IC 95% [−0,235; +0,153] y cambia de signo entre bloques. Se conserva como hipótesis secundaria, no como modelo promovido."
    )

    # 9.
    add_heading(doc, "9. Alerta de estrés: buen ranking, mala calibración", page_break=True)
    add_body(
        doc,
        "El evento de estrés se define dinámicamente como un cambio h=6 superior al percentil 80 de los cambios disponibles en train. El corte varía entre 0,169 y 0,292 pp. Entre los 46 orígenes hubo 16 positivos, concentrados en dos episodios efectivos."
    )
    add_table(
        doc,
        ["Modelo", "AP", "ROC-AUC", "Brier", "Precision", "Recall", "Alertas"],
        [
            ["Prevalencia", "0,339", "0,475", "0,248", "0,0%", "0,0%", "0,0%"],
            ["Logística", "0,771", "0,788", "0,277", "60,0%", "56,3%", "32,6%"],
            ["Random Forest", "0,466", "0,673", "0,254", "41,0%", "100,0%", "84,8%"],
            ["XGBoost", "0,460", "0,642", "0,299", "48,0%", "75,0%", "54,3%"],
        ],
        [2000, 1000, 1200, 1100, 1300, 1200, 1560],
        "reports/tables/stress_alert_metrics_development.csv.",
    )
    add_figure(
        doc,
        "stress_alert_evaluation.png",
        "Figura 4. Trayectoria y calibración de la alerta de estrés",
        "reports/figures/stress_alert_evaluation.png.",
        width=6.4,
    )
    add_body(
        doc,
        "La logística es el mejor ranking, pero asigna probabilidades casi binarias y su Brier es peor que la prevalencia histórica. Detecta el episodio septiembre 2021–noviembre 2022 con seis meses de retraso y omite agosto de 2023. Random Forest logra recall total al costo de 23 falsas alarmas."
    )
    add_callout(
        doc,
        "Decisión",
        "Ningún clasificador se aprueba como probabilidad calibrada ni como alerta operativa. La logística queda como challenger de ranking para investigación y revisión humana.",
        tone="red",
    )

    # 10.
    add_heading(doc, "10. Diagnóstico por regímenes", page_break=True)
    add_body(
        doc,
        "Los cortes por pandemia, inflación, TPM y actividad se aplican exclusivamente para describir heterogeneidad. Pueden solaparse y no intervienen en la selección del modelo."
    )
    add_figure(
        doc,
        "regime_performance.png",
        "Figura 5. Desempeño fuera de muestra por régimen macrofinanciero",
        "reports/figures/regime_performance.png.",
        width=6.4,
    )
    add_table(
        doc,
        ["Régimen", "Regresión h=6", "Alerta logística"],
        [
            ["Inflación alta", "ElasticNet mejora 6,5%", "AP 0,861; recall 64,3%"],
            ["Inflación baja", "ElasticNet empeora 40,0%", "AP 0,208; recall 0,0%"],
            ["TPM alta", "ElasticNet empeora 15,3%", "AP 0,925; recall 81,8%"],
            ["TPM baja", "ElasticNet empeora 17,9%", "AP 0,363; recall 0,0%"],
            ["Contracción", "ElasticNet empeora 6,0%", "AP 1,000; solo 2 eventos"],
        ],
        [2600, 3000, 3760],
        "reports/tables/regression_metrics_by_regime.csv y stress_alert_metrics_by_regime.csv.",
    )
    add_body(
        doc,
        "La aparente ventaja de ElasticNet con inflación alta y de la alerta bajo TPM alta es descriptiva. El tamaño reducido de cada submuestra, la concentración de eventos y la multiplicidad de cortes impiden convertir estos patrones en reglas de despliegue."
    )

    # 11.
    add_heading(doc, "11. Conclusiones e implicancias", page_break=True)
    add_heading(doc, "Respuesta a la pregunta", level=2)
    add_body(
        doc,
        "Con los datos, el horizonte y el protocolo definidos, el conjunto macrofinanciero no mejora de manera robusta el pronóstico puntual h=6 frente a cambio cero. ElasticNet es el mejor modelo aprendido, pero su error es mayor y sus drivers son inestables. La evidencia a doce meses es sugerente, no concluyente."
    )
    add_heading(doc, "Decisiones finales", level=2)
    decisions = [
        ("Campeón de regresión", "Cambio cero; MAE 0,434 pp."),
        ("Challenger de regresión", "ElasticNet completo de 23 variables."),
        ("Hipótesis secundaria", "ElasticNet h=12 con ventana expansiva."),
        ("Challenger de estrés", "Regresión logística como ranking exploratorio."),
        ("Alerta operativa", "No aprobada por calibración, estabilidad y escasez de episodios."),
        ("Holdout", "Cerrado; no se reportan resultados finales."),
    ]
    add_table(doc, ["Decisión", "Estado"], [[a, b] for a, b in decisions], [2900, 6460], "Model card v1.")
    add_heading(doc, "Implicancias prácticas", level=2)
    for item in [
        "Mantener una regla simple como referencia operacional y exigir mejoras repetibles antes de reemplazarla.",
        "Usar los modelos aprendidos como instrumentos de diagnóstico, no como motores automáticos de decisión.",
        "Evaluar la utilidad de la alerta con una función de costos explícita antes de elegir recall o precisión como prioridad.",
        "Acumular nuevos meses y episodios antes de abrir el holdout o fijar reglas condicionales por régimen.",
    ]:
        add_list_item(doc, item)

    # 12.
    add_heading(doc, "12. Limitaciones y agenda de investigación", page_break=True)
    for item in [
        "La evaluación externa contiene 46 orígenes y solo dos episodios efectivos de estrés; la precisión inferencial es limitada.",
        "Los targets consecutivos se solapan por el horizonte de seis meses; el bootstrap por bloques reduce, pero no elimina, esta dificultad.",
        "Las series macroeconómicas son de último vintage y no reconstruyen todas las revisiones históricas visibles en cada fecha.",
        "La pandemia, el ciclo inflacionario y cambios regulatorios generan quiebres que pueden dominar una muestra corta.",
        "El análisis agregado no representa diferencias entre bancos, productos, cohortes ni composición de cartera.",
        "No se estiman relaciones causales ni se cuantifican costos económicos de falsas alertas y omisiones.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "Siguientes extensiones", level=2)
    for index, item in enumerate([
        "Incorporar vintages históricos o snapshots mensuales de las series macro.",
        "Reevaluar h=12 cuando existan nuevos orígenes completamente fuera de muestra.",
        "Probar calibración temporal de probabilidades solo con una base más larga y episodios adicionales.",
        "Extender a panel por banco si la disponibilidad y comparabilidad institucional lo permiten.",
        "Definir una función de pérdida económica para la alerta antes de seleccionar un umbral operativo.",
    ]):
        add_list_item(doc, item, numbered=True, restart=index == 0)

    # 13.
    add_heading(doc, "13. Reproducibilidad y entrega", page_break=True)
    add_body(
        doc,
        "El repositorio separa datos raw, capas intermedias, tablas, figuras, configuración y código. Cada descarga cruda conserva hash SHA-256 y el proceso local reconstruye determinísticamente el análisis sin acceso a internet."
    )
    add_callout(
        doc,
        "Estado de aceptación",
        "11 de 11 controles pasan: cobertura, disponibilidad, feature count, filas de desarrollo, holdout reservado, orígenes de regresión y alerta, límites de probabilidad, artefactos requeridos y exclusión de .env.",
    )
    add_heading(doc, "Ejecución completa", level=2)
    add_body(doc, "Desde la raíz del repositorio, con el entorno activado y los archivos raw disponibles:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("$env:PYTHONPATH = \"src\"\npython -m morosidad_bancaria run-all-local\npython -m unittest discover -s tests -v"), size=9.5, color=NAVY)
    add_heading(doc, "Artefactos de trazabilidad", level=2)
    for item in [
        "data/metadata/technical_acceptance_v001.json: controles de aceptación.",
        "data/metadata/reproduction_manifest_v001.json: entorno y hashes de archivos.",
        "docs/model_card_v1.md: alcance, uso previsto y estado de los modelos.",
        "docs/reproduction.md: instalación, ejecución y verificación.",
        "reports/tables/ y reports/figures/: evidencia tabular y visual de desarrollo.",
    ]:
        add_list_item(doc, item)
    add_callout(
        doc,
        "Seguridad",
        "Las credenciales residen únicamente en .env, que está excluido de Git. El programa no las incorpora en rutas, manifiestos ni logs. El informe no contiene usuarios, contraseñas ni tokens.",
        tone="gold",
    )

    # Referencias.
    add_heading(doc, "Referencias y archivos base", page_break=True)
    refs = [
        "[1] Comisión para el Mercado Financiero (CMF). APIBEST y comunicados estadísticos mensuales del sistema bancario chileno.",
        "[2] Banco Central de Chile. Base de Datos Estadísticos (BDE), series macroeconómicas mensuales.",
        "[3] docs/project_charter.md y docs/decisions/ADR-001 a ADR-013. Decisiones metodológicas del proyecto.",
        "[4] docs/baseline_backtest_v1.md, docs/ml_backtest_v1.md y docs/robustness_v1.md. Resultados de regresión.",
        "[5] docs/explainability_v1.md y docs/horizon_robustness_v1.md. Drivers y sensibilidad.",
        "[6] docs/stress_alert_v1.md y docs/technical_closure_v1.md. Alerta de estrés y regímenes.",
        "[7] docs/model_card_v1.md y docs/reproduction.md. Gobierno, limitaciones y reproducción.",
    ]
    for ref in refs:
        add_list_item(doc, ref, numbered=False)
    add_heading(doc, "Criterio de cierre", level=2)
    add_body(
        doc,
        "El hito 10 se considera concluido con la entrega conjunta de este informe, la presentación final, el manual de reproducción, la model card, las tablas y figuras auditables, y la verificación técnica sin acceso al holdout."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
